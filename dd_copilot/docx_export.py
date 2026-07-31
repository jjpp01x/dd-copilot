"""Export the Markdown report as a Word document.

Markdown is the working format; `.docx` is what a client actually receives.
The conversion is deliberately narrow — it understands only the constructs
`report.py` emits (headings, bullets, one pipe table, paragraphs) rather than
pulling in a general Markdown engine for a document whose shape we control.

The claims table is the reason this module exists: rendered as pipe-separated
text it reads as noise, and it is the section worth reading.
"""

from __future__ import annotations

import re
from pathlib import Path

HEADING_PREFIX = "#"
BULLET_PREFIX = "- "


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _is_separator_row(cells: list[str]) -> bool:
    return all(set(cell) <= set("-: ") and "-" in cell for cell in cells)


#: A cell boundary is a pipe that is not escaped. Splitting on every pipe would
#: tear apart cells whose own text contains one — which report.py escapes
#: precisely so the markdown table survives.
_CELL_BOUNDARY = re.compile(r"(?<!\\)\|")


def _split_row(line: str) -> list[str]:
    # Strip the outer pipes, split on unescaped ones, then unescape: Word wants
    # the original text, not report.py's markdown-safe form.
    return [cell.strip().replace("\\|", "|") for cell in _CELL_BOUNDARY.split(line[1:-1])]


def _add_table(document, rows: list[list[str]]) -> None:
    header, *body = rows
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, header):
        cell.text = text
    for row in body:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text


def write_docx(markdown: str, path: str | Path) -> Path:
    """Renders `markdown` to a .docx at `path` and returns the path."""
    from docx import Document

    document = Document()
    pending_table: list[list[str]] = []

    def flush_table() -> None:
        if pending_table:
            _add_table(document, pending_table)
            pending_table.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()

        if _is_table_row(line):
            cells = _split_row(line)
            if not _is_separator_row(cells):
                pending_table.append(cells)
            continue

        flush_table()

        if not line.strip():
            continue

        if line.startswith(HEADING_PREFIX):
            level = len(line) - len(line.lstrip(HEADING_PREFIX))
            text = line[level:].strip()
            # Markdown h1 is the report title; Word's Title style suits it
            # better than Heading 1, which is then free for the sections.
            if level == 1:
                document.add_heading(text, level=0)
            else:
                document.add_heading(text, level=min(level - 1, 9))
        elif line.startswith(BULLET_PREFIX):
            document.add_paragraph(line[len(BULLET_PREFIX):].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    flush_table()

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path
